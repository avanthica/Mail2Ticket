import os
from docx import Document as DocxDocument
from langchain_chroma import Chroma
from langchain_core.documents import Document
from langchain_google_genai import GoogleGenerativeAIEmbeddings

api_key = 'AIzaSyDThUorg68gXA6gN7ohTIcDFd3OVkWhk70'
persist_directory = 'chroma_db'
#Step 1: Parse .docx and Extract Email Chunks

def extract_email_chunks(docx_path):
    doc = DocxDocument(docx_path)
    email_chunks = []
    
    current_email = []
    for para in doc.paragraphs:
        if para.text.startswith("Mail"):
            if current_email:
                email_chunks.append("\n".join(current_email))
            current_email = [para.text]
        else:
            current_email.append(para.text)
    
    if current_email:
        email_chunks.append("\n".join(current_email))
    
    return email_chunks

def process_docx_folder(folder_path):
    all_chunks = {}
    for filename in os.listdir(folder_path):
        if filename.endswith(".docx"):
            docx_path = os.path.join(folder_path, filename)
            email_chunks = extract_email_chunks(docx_path)
            all_chunks[filename] = email_chunks
    return all_chunks

 #Step 2: Convert Chunks to Text + Metadata

def flatten_chunks(all_chunks):
    texts = []
    metadatas = []

    for filename, chunks in all_chunks.items():
        for i, chunk in enumerate(chunks):
            texts.append(chunk)
            metadatas.append({
                "source": filename,
                "mail_id": f"Mail {i+1}"
            })
    return texts, metadatas

def get_vector_store(texts, metadatas, api_key, persist_directory):
    embeddings = GoogleGenerativeAIEmbeddings(
        model="models/embedding-001",
        google_api_key=api_key
    )
    
    # Create Document objects with metadata
    documents = [Document(page_content=texts[i], metadata=metadatas[i]) for i in range(len(texts))]

    vectordb = Chroma.from_documents(
        documents=documents,
        embedding=embeddings,
        persist_directory=persist_directory
    )
    # vectordb.persist()
    # vectordb = Chroma(
    # persist_directory=persist_directory,
    # embedding_function=embeddings
    # )
    print("📦 Total stored chunks:", len(vectordb.get()['ids']))
    return vectordb

folder_path = "convo mails"

all_chunks = process_docx_folder(folder_path)
texts, metadatas = flatten_chunks(all_chunks)
vectordb = get_vector_store(texts,metadatas,api_key,persist_directory)
